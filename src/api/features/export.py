"""
Export functionality for Parallax Pal

Provides multiple export formats for research results including
PDF, DOCX, Notion, and custom integrations.
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, 
    Table, TableStyle, PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart

import pandas as pd
import json
from typing import Dict, List, Any, Optional
from datetime import datetime
import io
import base64
import asyncio
import logging
from pathlib import Path

# Additional export libraries
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from notion_client import AsyncClient as NotionClient
import csv
import xlsxwriter

logger = logging.getLogger(__name__)


class ResearchExporter:
    """Export research results in multiple formats"""
    
    def __init__(self):
        """Initialize exporter with styles and templates"""
        
        # PDF styles
        self.pdf_styles = self._create_pdf_styles()
        
        # Export templates
        self.templates = {
            'academic': self._get_academic_template(),
            'business': self._get_business_template(),
            'casual': self._get_casual_template()
        }
        
        logger.info("Research exporter initialized")
    
    async def export_research(
        self,
        research_data: Dict[str, Any],
        format: str,
        template: str = "academic",
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Export research in specified format
        
        Args:
            research_data: Research results to export
            format: Export format (pdf, docx, txt, json, csv, notion)
            template: Template style
            options: Format-specific options
            
        Returns:
            Export result with file data or URL
        """
        options = options or {}
        
        exporters = {
            'pdf': self.export_to_pdf,
            'docx': self.export_to_docx,
            'txt': self.export_to_text,
            'json': self.export_to_json,
            'csv': self.export_to_csv,
            'xlsx': self.export_to_excel,
            'notion': self.export_to_notion,
            'markdown': self.export_to_markdown,
            'html': self.export_to_html
        }
        
        exporter = exporters.get(format)
        if not exporter:
            raise ValueError(f"Unsupported export format: {format}")
        
        try:
            result = await exporter(research_data, template, options)
            
            # Log export
            logger.info(f"Research exported to {format} format")
            
            return {
                'success': True,
                'format': format,
                'template': template,
                **result
            }
            
        except Exception as e:
            logger.error(f"Export error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def export_to_pdf(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as PDF report"""
        
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4 if options.get('pagesize') == 'A4' else letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Build story
        story = []
        styles = self.pdf_styles
        
        # Title page
        story.append(Paragraph(research_data.get('title', 'Research Report'), styles['CustomTitle']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y')}",
            styles['DateStyle']
        ))
        story.append(Spacer(1, 0.5*inch))
        
        # Query
        story.append(Paragraph("Research Query", styles['Heading1']))
        story.append(Paragraph(research_data.get('query', ''), styles['QueryStyle']))
        story.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        if 'summary' in research_data:
            story.append(Paragraph("Executive Summary", styles['Heading1']))
            story.append(Paragraph(research_data['summary'], styles['BodyText']))
            story.append(Spacer(1, 0.3*inch))
        
        # Key Findings
        if 'findings' in research_data:
            story.append(Paragraph("Key Findings", styles['Heading1']))
            for i, finding in enumerate(research_data['findings'], 1):
                story.append(Paragraph(f"{i}. {finding}", styles['FindingStyle']))
            story.append(Spacer(1, 0.3*inch))
        
        # Detailed Analysis
        if 'analysis' in research_data:
            story.append(PageBreak())
            story.append(Paragraph("Detailed Analysis", styles['Heading1']))
            
            for section in research_data['analysis']:
                story.append(Paragraph(section.get('title', ''), styles['Heading2']))
                story.append(Paragraph(section.get('content', ''), styles['BodyText']))
                story.append(Spacer(1, 0.2*inch))
        
        # Sources
        if 'sources' in research_data:
            story.append(PageBreak())
            story.append(Paragraph("Sources", styles['Heading1']))
            
            # Create source table
            source_data = [['Title', 'URL', 'Reliability']]
            for source in research_data['sources']:
                source_data.append([
                    source.get('title', '')[:50] + '...' if len(source.get('title', '')) > 50 else source.get('title', ''),
                    source.get('url', '')[:40] + '...' if len(source.get('url', '')) > 40 else source.get('url', ''),
                    f"{source.get('reliability', 0) * 100:.0f}%"
                ])
            
            source_table = Table(source_data, colWidths=[3*inch, 2.5*inch, 1*inch])
            source_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(source_table)
        
        # Citations
        if 'citations' in research_data:
            story.append(PageBreak())
            story.append(Paragraph("Bibliography", styles['Heading1']))
            
            for citation in research_data['citations']:
                story.append(Paragraph(citation, styles['Citation']))
                story.append(Spacer(1, 0.1*inch))
        
        # Knowledge Graph Summary
        if 'knowledge_graph' in research_data:
            story.append(PageBreak())
            story.append(Paragraph("Knowledge Graph Summary", styles['Heading1']))
            
            kg = research_data['knowledge_graph']
            story.append(Paragraph(
                f"The knowledge graph contains {len(kg.get('nodes', []))} entities "
                f"and {len(kg.get('edges', []))} relationships.",
                styles['BodyText']
            ))
            
            # Top entities
            if kg.get('nodes'):
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph("Key Entities:", styles['Heading2']))
                
                entity_types = {}
                for node in kg['nodes'][:20]:  # Top 20 entities
                    node_type = node.get('type', 'Unknown')
                    if node_type not in entity_types:
                        entity_types[node_type] = []
                    entity_types[node_type].append(node.get('label', ''))
                
                for entity_type, entities in entity_types.items():
                    story.append(Paragraph(f"<b>{entity_type}:</b> {', '.join(entities)}", styles['BodyText']))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF data
        buffer.seek(0)
        pdf_data = buffer.getvalue()
        
        return {
            'data': base64.b64encode(pdf_data).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            'size': len(pdf_data)
        }
    
    async def export_to_docx(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as Word document"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(research_data.get('title', 'Research Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Template: {template.capitalize()}")
        
        # Add query
        doc.add_heading('Research Query', level=1)
        query_para = doc.add_paragraph(research_data.get('query', ''))
        query_para.style = 'Quote'
        
        # Add summary
        if 'summary' in research_data:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(research_data['summary'])
        
        # Add findings
        if 'findings' in research_data:
            doc.add_heading('Key Findings', level=1)
            for finding in research_data['findings']:
                doc.add_paragraph(finding, style='List Bullet')
        
        # Add detailed analysis
        if 'analysis' in research_data:
            doc.add_page_break()
            doc.add_heading('Detailed Analysis', level=1)
            
            for section in research_data['analysis']:
                doc.add_heading(section.get('title', ''), level=2)
                doc.add_paragraph(section.get('content', ''))
        
        # Add sources table
        if 'sources' in research_data:
            doc.add_page_break()
            doc.add_heading('Sources', level=1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Title'
            header_cells[1].text = 'URL'
            header_cells[2].text = 'Reliability'
            
            # Data rows
            for source in research_data['sources']:
                row_cells = table.add_row().cells
                row_cells[0].text = source.get('title', '')
                row_cells[1].text = source.get('url', '')
                row_cells[2].text = f"{source.get('reliability', 0) * 100:.0f}%"
        
        # Add citations
        if 'citations' in research_data:
            doc.add_page_break()
            doc.add_heading('Bibliography', level=1)
            
            for citation in research_data['citations']:
                doc.add_paragraph(citation, style='Bibliography')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_data = buffer.getvalue()
        
        return {
            'data': base64.b64encode(docx_data).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            'size': len(docx_data)
        }
    
    async def export_to_notion(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research to Notion"""
        
        notion_token = options.get('notion_token')
        parent_page_id = options.get('parent_page_id')
        
        if not notion_token:
            raise ValueError("Notion integration token required")
        
        notion = NotionClient(auth=notion_token)
        
        # Create page properties
        properties = {
            "title": {
                "title": [{"text": {"content": research_data.get('title', 'Research Report')}}]
            },
            "Date": {
                "date": {"start": datetime.now().isoformat()}
            },
            "Query": {
                "rich_text": [{"text": {"content": research_data.get('query', '')}}]
            }
        }
        
        # Create page children (content blocks)
        children = []
        
        # Add summary
        if 'summary' in research_data:
            children.append({
                "object": "block",
                "type": "heading_1",
                "heading_1": {
                    "rich_text": [{"text": {"content": "Executive Summary"}}]
                }
            })
            children.append({
                "object": "block",
                "type": "paragraph",
                "paragraph": {
                    "rich_text": [{"text": {"content": research_data['summary']}}]
                }
            })
        
        # Add findings
        if 'findings' in research_data:
            children.append({
                "object": "block",
                "type": "heading_1",
                "heading_1": {
                    "rich_text": [{"text": {"content": "Key Findings"}}]
                }
            })
            
            for finding in research_data['findings']:
                children.append({
                    "object": "block",
                    "type": "bulleted_list_item",
                    "bulleted_list_item": {
                        "rich_text": [{"text": {"content": finding}}]
                    }
                })
        
        # Add sources
        if 'sources' in research_data:
            children.append({
                "object": "block",
                "type": "heading_1",
                "heading_1": {
                    "rich_text": [{"text": {"content": "Sources"}}]
                }
            })
            
            # Create table
            source_rows = []
            for source in research_data['sources']:
                source_rows.append([
                    source.get('title', ''),
                    source.get('url', ''),
                    f"{source.get('reliability', 0) * 100:.0f}%"
                ])
            
            children.append({
                "object": "block",
                "type": "table",
                "table": {
                    "table_width": 3,
                    "has_column_header": True,
                    "has_row_header": False,
                    "children": [
                        {
                            "object": "block",
                            "type": "table_row",
                            "table_row": {
                                "cells": [
                                    [{"text": {"content": "Title"}}],
                                    [{"text": {"content": "URL"}}],
                                    [{"text": {"content": "Reliability"}}]
                                ]
                            }
                        }
                    ] + [
                        {
                            "object": "block",
                            "type": "table_row",
                            "table_row": {
                                "cells": [
                                    [{"text": {"content": cell}}] for cell in row
                                ]
                            }
                        } for row in source_rows
                    ]
                }
            })
        
        # Create the page
        parent = {"page_id": parent_page_id} if parent_page_id else {"type": "workspace"}
        
        response = await notion.pages.create(
            parent=parent,
            properties=properties,
            children=children
        )
        
        return {
            'notion_url': response['url'],
            'page_id': response['id']
        }
    
    async def export_to_markdown(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as Markdown"""
        
        md_lines = []
        
        # Title
        md_lines.append(f"# {research_data.get('title', 'Research Report')}")
        md_lines.append("")
        md_lines.append(f"*Generated on {datetime.now().strftime('%B %d, %Y')}*")
        md_lines.append("")
        
        # Query
        md_lines.append("## Research Query")
        md_lines.append(f"> {research_data.get('query', '')}")
        md_lines.append("")
        
        # Summary
        if 'summary' in research_data:
            md_lines.append("## Executive Summary")
            md_lines.append(research_data['summary'])
            md_lines.append("")
        
        # Findings
        if 'findings' in research_data:
            md_lines.append("## Key Findings")
            for finding in research_data['findings']:
                md_lines.append(f"- {finding}")
            md_lines.append("")
        
        # Analysis
        if 'analysis' in research_data:
            md_lines.append("## Detailed Analysis")
            for section in research_data['analysis']:
                md_lines.append(f"### {section.get('title', '')}")
                md_lines.append(section.get('content', ''))
                md_lines.append("")
        
        # Sources
        if 'sources' in research_data:
            md_lines.append("## Sources")
            md_lines.append("")
            md_lines.append("| Title | URL | Reliability |")
            md_lines.append("|-------|-----|-------------|")
            
            for source in research_data['sources']:
                title = source.get('title', '')
                url = source.get('url', '')
                reliability = f"{source.get('reliability', 0) * 100:.0f}%"
                md_lines.append(f"| {title} | [{url}]({url}) | {reliability} |")
            md_lines.append("")
        
        # Knowledge Graph
        if 'knowledge_graph' in research_data:
            kg = research_data['knowledge_graph']
            md_lines.append("## Knowledge Graph Summary")
            md_lines.append(f"- **Entities**: {len(kg.get('nodes', []))}")
            md_lines.append(f"- **Relationships**: {len(kg.get('edges', []))}")
            md_lines.append("")
            
            # Entity breakdown
            entity_types = {}
            for node in kg.get('nodes', []):
                node_type = node.get('type', 'Unknown')
                entity_types[node_type] = entity_types.get(node_type, 0) + 1
            
            if entity_types:
                md_lines.append("### Entity Types")
                for entity_type, count in entity_types.items():
                    md_lines.append(f"- **{entity_type}**: {count}")
                md_lines.append("")
        
        # Citations
        if 'citations' in research_data:
            md_lines.append("## Bibliography")
            for i, citation in enumerate(research_data['citations'], 1):
                md_lines.append(f"{i}. {citation}")
            md_lines.append("")
        
        markdown_content = '\n'.join(md_lines)
        
        return {
            'data': base64.b64encode(markdown_content.encode('utf-8')).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
            'size': len(markdown_content)
        }
    
    async def export_to_excel(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research data to Excel with multiple sheets"""
        
        buffer = io.BytesIO()
        
        with xlsxwriter.Workbook(buffer) as workbook:
            # Formats
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#D7E4BD',
                'border': 1
            })
            
            # Summary sheet
            summary_sheet = workbook.add_worksheet('Summary')
            summary_sheet.write('A1', research_data.get('title', 'Research Report'), title_format)
            summary_sheet.write('A3', 'Query:')
            summary_sheet.write('B3', research_data.get('query', ''))
            summary_sheet.write('A4', 'Generated:')
            summary_sheet.write('B4', datetime.now().strftime('%Y-%m-%d %H:%M'))
            
            if 'summary' in research_data:
                summary_sheet.write('A6', 'Executive Summary:')
                summary_sheet.write('A7', research_data['summary'])
            
            # Findings sheet
            if 'findings' in research_data:
                findings_sheet = workbook.add_worksheet('Key Findings')
                findings_sheet.write('A1', 'Key Findings', title_format)
                
                for i, finding in enumerate(research_data['findings'], 3):
                    findings_sheet.write(f'A{i}', f'{i-2}. {finding}')
            
            # Sources sheet
            if 'sources' in research_data:
                sources_sheet = workbook.add_worksheet('Sources')
                
                # Headers
                headers = ['Title', 'URL', 'Reliability', 'Summary']
                for col, header in enumerate(headers):
                    sources_sheet.write(0, col, header, header_format)
                
                # Data
                for row, source in enumerate(research_data['sources'], 1):
                    sources_sheet.write(row, 0, source.get('title', ''))
                    sources_sheet.write(row, 1, source.get('url', ''))
                    sources_sheet.write(row, 2, source.get('reliability', 0))
                    sources_sheet.write(row, 3, source.get('summary', ''))
                
                # Adjust column widths
                sources_sheet.set_column('A:A', 40)
                sources_sheet.set_column('B:B', 50)
                sources_sheet.set_column('C:C', 12)
                sources_sheet.set_column('D:D', 60)
            
            # Knowledge Graph sheet
            if 'knowledge_graph' in research_data:
                kg = research_data['knowledge_graph']
                
                # Nodes sheet
                nodes_sheet = workbook.add_worksheet('KG Nodes')
                node_headers = ['ID', 'Label', 'Type', 'Properties']
                for col, header in enumerate(node_headers):
                    nodes_sheet.write(0, col, header, header_format)
                
                for row, node in enumerate(kg.get('nodes', []), 1):
                    nodes_sheet.write(row, 0, node.get('id', ''))
                    nodes_sheet.write(row, 1, node.get('label', ''))
                    nodes_sheet.write(row, 2, node.get('type', ''))
                    nodes_sheet.write(row, 3, json.dumps(node.get('properties', {})))
                
                # Edges sheet
                edges_sheet = workbook.add_worksheet('KG Edges')
                edge_headers = ['Source', 'Target', 'Type', 'Weight']
                for col, header in enumerate(edge_headers):
                    edges_sheet.write(0, col, header, header_format)
                
                for row, edge in enumerate(kg.get('edges', []), 1):
                    edges_sheet.write(row, 0, edge.get('source', ''))
                    edges_sheet.write(row, 1, edge.get('target', ''))
                    edges_sheet.write(row, 2, edge.get('type', ''))
                    edges_sheet.write(row, 3, edge.get('weight', 0))
        
        buffer.seek(0)
        excel_data = buffer.getvalue()
        
        return {
            'data': base64.b64encode(excel_data).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            'size': len(excel_data)
        }
    
    async def export_to_text(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as plain text"""
        
        lines = []
        
        # Title
        lines.append(research_data.get('title', 'Research Report').upper())
        lines.append('=' * len(research_data.get('title', 'Research Report')))
        lines.append('')
        lines.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        lines.append('')
        
        # Query
        lines.append('RESEARCH QUERY')
        lines.append('-' * 14)
        lines.append(research_data.get('query', ''))
        lines.append('')
        
        # Summary
        if 'summary' in research_data:
            lines.append('EXECUTIVE SUMMARY')
            lines.append('-' * 17)
            lines.append(research_data['summary'])
            lines.append('')
        
        # Findings
        if 'findings' in research_data:
            lines.append('KEY FINDINGS')
            lines.append('-' * 12)
            for i, finding in enumerate(research_data['findings'], 1):
                lines.append(f"{i}. {finding}")
            lines.append('')
        
        # Sources
        if 'sources' in research_data:
            lines.append('SOURCES')
            lines.append('-' * 7)
            for i, source in enumerate(research_data['sources'], 1):
                lines.append(f"{i}. {source.get('title', '')}")
                lines.append(f"   URL: {source.get('url', '')}")
                lines.append(f"   Reliability: {source.get('reliability', 0) * 100:.0f}%")
                lines.append('')
        
        text_content = '\n'.join(lines)
        
        return {
            'data': base64.b64encode(text_content.encode('utf-8')).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            'size': len(text_content)
        }
    
    async def export_to_json(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as JSON"""
        
        # Add metadata
        export_data = {
            'metadata': {
                'title': research_data.get('title', 'Research Report'),
                'generated_at': datetime.now().isoformat(),
                'template': template,
                'version': '1.0'
            },
            'research': research_data
        }
        
        json_content = json.dumps(export_data, indent=2, ensure_ascii=False)
        
        return {
            'data': base64.b64encode(json_content.encode('utf-8')).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            'size': len(json_content)
        }
    
    async def export_to_csv(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research sources as CSV"""
        
        buffer = io.StringIO()
        
        # Define fields
        fieldnames = ['title', 'url', 'reliability', 'summary', 'type']
        writer = csv.DictWriter(buffer, fieldnames=fieldnames)
        
        # Write header
        writer.writeheader()
        
        # Write sources
        for source in research_data.get('sources', []):
            writer.writerow({
                'title': source.get('title', ''),
                'url': source.get('url', ''),
                'reliability': source.get('reliability', 0),
                'summary': source.get('summary', ''),
                'type': source.get('type', 'web')
            })
        
        csv_content = buffer.getvalue()
        
        return {
            'data': base64.b64encode(csv_content.encode('utf-8')).decode('utf-8'),
            'filename': f"research_sources_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            'size': len(csv_content)
        }
    
    async def export_to_html(
        self,
        research_data: Dict[str, Any],
        template: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export research as HTML with embedded styles"""
        
        html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        .query {{ background-color: #ecf0f1; padding: 15px; border-left: 4px solid #3498db; margin: 20px 0; }}
        .finding {{ background-color: #e8f5e9; padding: 10px; margin: 10px 0; border-radius: 5px; }}
        .source {{ border: 1px solid #ddd; padding: 10px; margin: 10px 0; border-radius: 5px; }}
        .reliability {{ color: #27ae60; font-weight: bold; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #3498db; color: white; }}
        .metadata {{ color: #7f8c8d; font-size: 0.9em; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="metadata">Generated on {date}</p>
    
    <div class="query">
        <h2>Research Query</h2>
        <p>{query}</p>
    </div>
    
    {content}
    
    <footer>
        <p class="metadata">Generated by Parallax Pal Research Assistant</p>
    </footer>
</body>
</html>"""
        
        # Build content sections
        content_parts = []
        
        # Summary
        if 'summary' in research_data:
            content_parts.append(f"""
    <section>
        <h2>Executive Summary</h2>
        <p>{research_data['summary']}</p>
    </section>""")
        
        # Findings
        if 'findings' in research_data:
            findings_html = '\n'.join([
                f'<div class="finding">{finding}</div>'
                for finding in research_data['findings']
            ])
            content_parts.append(f"""
    <section>
        <h2>Key Findings</h2>
        {findings_html}
    </section>""")
        
        # Sources
        if 'sources' in research_data:
            sources_html = '\n'.join([
                f"""<div class="source">
                    <h3>{source.get('title', '')}</h3>
                    <p><a href="{source.get('url', '')}" target="_blank">{source.get('url', '')}</a></p>
                    <p class="reliability">Reliability: {source.get('reliability', 0) * 100:.0f}%</p>
                </div>"""
                for source in research_data['sources']
            ])
            content_parts.append(f"""
    <section>
        <h2>Sources</h2>
        {sources_html}
    </section>""")
        
        # Format HTML
        html_content = html_template.format(
            title=research_data.get('title', 'Research Report'),
            date=datetime.now().strftime('%B %d, %Y'),
            query=research_data.get('query', ''),
            content='\n'.join(content_parts)
        )
        
        return {
            'data': base64.b64encode(html_content.encode('utf-8')).decode('utf-8'),
            'filename': f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
            'size': len(html_content)
        }
    
    def _create_pdf_styles(self) -> Dict[str, Any]:
        """Create PDF paragraph styles"""
        
        styles = getSampleStyleSheet()
        
        # Custom title style
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        # Date style
        styles.add(ParagraphStyle(
            name='DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#7f8c8d'),
            alignment=TA_CENTER
        ))
        
        # Query style
        styles.add(ParagraphStyle(
            name='QueryStyle',
            parent=styles['Normal'],
            fontSize=12,
            leftIndent=20,
            rightIndent=20,
            textColor=colors.HexColor('#34495e'),
            backColor=colors.HexColor('#ecf0f1'),
            borderColor=colors.HexColor('#3498db'),
            borderWidth=2,
            borderPadding=10,
            alignment=TA_JUSTIFY
        ))
        
        # Finding style
        styles.add(ParagraphStyle(
            name='FindingStyle',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=10,
            textColor=colors.HexColor('#27ae60')
        ))
        
        # Citation style
        styles.add(ParagraphStyle(
            name='Citation',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            textColor=colors.HexColor('#7f8c8d')
        ))
        
        return styles
    
    def _get_academic_template(self) -> Dict[str, Any]:
        """Get academic report template settings"""
        return {
            'font': 'Times-Roman',
            'heading_style': 'formal',
            'include_toc': True,
            'citation_style': 'APA',
            'page_numbers': True
        }
    
    def _get_business_template(self) -> Dict[str, Any]:
        """Get business report template settings"""
        return {
            'font': 'Helvetica',
            'heading_style': 'modern',
            'include_toc': False,
            'citation_style': 'simple',
            'page_numbers': True,
            'include_executive_summary': True
        }
    
    def _get_casual_template(self) -> Dict[str, Any]:
        """Get casual report template settings"""
        return {
            'font': 'Helvetica',
            'heading_style': 'simple',
            'include_toc': False,
            'citation_style': 'minimal',
            'page_numbers': False
        }